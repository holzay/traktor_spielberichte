import requests
from bs4 import BeautifulSoup
from func import extract_date,traktor_clean_article
import os
from docx import Document
from docx.shared import Inches

header = {
	'User-Agent' : 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:97.0) Gecko/20100101 Firefox/97.0'
}


chapter_dict = {

'Verbandsliga 2011/2012, Kleinfeld' : '2011-2012_Verbandsliga_Kleinfeld',
'Saison 2010/2011' : '2010-2011_Saison',
'Saison 2009/2010' : '2009-2010_Saison',
'Saison 2008/2009' : '2008-2009_Saison',
'Saison 2007/2008' : '2007-2008_Saison',
}

target_folder = 'export_old_mannschaft_1'

doc_name = 'Spielberichte_alt_Mannschaft_I.docx'

document = Document()

document.add_heading('Spielberichte Mannschaft I', 0)



if not os.path.isdir(f'./{target_folder}'):
		os.mkdir(target_folder)

os.chdir(r'D:\Code\traktor_spielberichte\export_old_mannschaft_1')


url = 'http://s600778792.online.de/i-mannschaft-vergangene-saisons'

r = requests.get(url,headers=header).text

soup = BeautifulSoup(r,'html.parser')

with open('output.html','w',encoding='utf8') as log:
	log.write(soup.prettify())

articles = soup.find_all('div',class_='sp-wrap sp-wrap-default')

article_id = 0
arcticle_dict = {}

for article in articles:
	chapter = article.find('div',class_='sp-head unfolded') or article.div.text
	chapter = chapter.replace('\n','').strip()
	document.add_heading(chapter, level=1)
	# if chapter == 'Kreisliga A | Saison 2012/2013':
	# 	print('debug!')
	folder_name = chapter_dict[chapter]
	if not os.path.isdir(f'./{folder_name}'):
		os.mkdir(folder_name)
	paragraphs = article.find_all('p')
	# file_open = False
	for p in paragraphs:
		if p.text.strip() == '':
			continue
		para = document.add_paragraph(p.text.strip())
		document.add_page_break()
		# strong_tag = p.find('strong')
		# multiple_strong_tags = p.find_all('strong')
		# if file_open:
		# 	f.close()
		# 	file_open = False
		# title = strong_tag.text.strip()
		# if not title[:1].isnumeric():
		# 	continue
		# article_id+=1
		# article_text = []
		# try:
		# 	date = extract_date(title,title)
		# except:
		# 	print(f'problem with {title}')
		# if len(multiple_strong_tags)>1:
		# 	for index,s_tag in enumerate(multiple_strong_tags):
		# 		if index == 0:
		# 			pass
		# 		else:
		# 			title = f'{title}\n{s_tag.text.strip()}'
		# f = open(f'../{target_folder}/{folder_name}/{article_id}.txt','w',encoding='utf8')
		
		# f.write(f'{title}\n')
		# file_open = True
		# texts = p.find_next_siblings('p')
	# if file_open:
	# 	if p.text.strip() != '':
	# 		f.write(f'{p.text.strip()}\n')
			
document.save(doc_name)
