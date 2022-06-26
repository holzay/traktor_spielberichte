from cgitb import text
import requests
from bs4 import BeautifulSoup
from func import extract_date,traktor_clean_article
import os
from docx import Document
from docx.shared import Inches

header = {
	'User-Agent' : 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:97.0) Gecko/20100101 Firefox/97.0'
}

doc_name = 'Spielberichte_Traktor_GF.docx'

document = Document()

document.add_heading('Spielberichte Traktor Großfeld alt', 0)


chapter_dict = {

'Fz-Verbandsliga | Saison 2016/17' : '2016-2017_FZ_Verbandsliga',
'Fz-Verbandsliga | Saison 2015/16' : '2015-2016_FZ_Verbandsliga',
'Landesliga I | Saison 2014/15' : '2014-2015_Landesliga_I',
'Bezirksliga II | Saison 2013/2014' : '2013-2014_Bezirksliga_II',
'Kreisliga A | Saison 2012/2013' : '2012-2013_Kreisliga_A'
}

os.chdir(r'D:\Code\traktor_spielberichte\export_gf_old')


url = 'http://s600778792.online.de/grossfeld_saisons'

r = requests.get(url,headers=header).text

soup = BeautifulSoup(r,'html.parser')


with open('output.html','w',encoding='utf8') as log:
	log.write(soup.prettify())

articles = soup.find_all('div',class_='sp-wrap sp-wrap-default')

article_id = 0
article_dict = {}

last_para = ''

broken_titles = ['7.6.15 Solidarität/Tasmania 0:2 (0:1)']

text_cont = []

switch_mode = False

for article in articles:
	chapter = article.find('div',class_='sp-head unfolded') or article.div.text
	chapter = chapter.replace('\n','').strip()
	document.add_heading(chapter, level=1)
	folder_name = chapter_dict[chapter]
	if not os.path.isdir(f'./{folder_name}'):
		os.mkdir(folder_name)
	paragraphs = article.find_all('p')
	for para in paragraphs:
		new_article = True
		if para.get_text() == '15.6.15 Rotation Prenzlauer Berg – SV Traktor Boxhagen 1:3 (0:0)':
			switch_mode = True
		# text_part = para.get_text().strip()
		spielbericht = para.find('span')
		if spielbericht:
			spielbericht = para.span.extract()
			text_part = para.get_text().strip()
			if text_part and not switch_mode:
				article_dict[article_id]['text'].append(text_part.strip())
				new_article = False
				# para.text.replace(text_part,'')
			# if article_id > 0 and text_part != '':
			# 	article_dict[article_id]['text'].append(text_part.strip())
				# text_part = para.get_text().strip().replace(text_part,'')
				# text_part.string.extract()
			article_id+=1
			title = spielbericht.text.strip()
			if title.startswith('31.1.2016'):
				print('debug')
			try:
				date,new_title,date_sort = extract_date(title)
			except:
				date,new_title,date_sort = 'unknown',title,'0'
			heading = f'{date} | {new_title}'
			article_dict[article_id] = {
				'date' : date,
				'title' : title,
				# 'link' : link,
				'date_sort' :  int(date_sort),
				'heading' : heading,
				'text' : []
			}
		text_part = para.get_text().strip()
		if (new_article or switch_mode) and article_id > 0 and text_part.strip() != '' and text_part.strip() != '[collapse]' and text_part != title:
			if text_part.startswith(title):
				text_part = text_part.replace(title,'')
			article_dict[article_id]['text'].append(text_part.strip())

			
	
	# file_open = False
	# for p in paragraphs:
	# 	text_part = p.text
	# 	strong_tag = p.find('strong')
	# 	multiple_strong_tags = p.find_all('strong')
	# 	if strong_tag:
	# 		if text_part !='' and text_part != strong_tag.text:
	# 			article_dict[article_id]['text'].append(text_part.strip())
	# 			p.extract()
	# 		article_id+=1
	# 		title = strong_tag.text.strip()
	# 		if not title[:1].isnumeric():
	# 			continue
	# 		try:
	# 			date,new_title,date_sort = extract_date(title)
	# 		except:
	# 			date,new_title,date_sort = 'unknown',title,'0'
	# 		if len(multiple_strong_tags)>1:
	# 			for index,s_tag in enumerate(multiple_strong_tags):
	# 				if index == 0:
	# 					pass
	# 				else:
	# 					title = f'{title}\n{s_tag.text.strip()}'
	# 		if date_sort == '20160131':
	# 			print('debug!')
	# 		heading = f'{date} | {new_title}'
	# 		article_dict[article_id] = {
	# 			'date' : date,
	# 			'title' : title,
	# 			# 'link' : link,
	# 			'date_sort' :  int(date_sort),
	# 			'heading' : heading,
	# 			'text' : []
	# 		}
	# 		p.strong.decompose()
	# 	if article_id !=0 and text_part:
	# 		if text_part.strip() != '' and text_part.strip() != '[collapse]':
	# 			if text_part == title:
	# 				continue
	# 			if text_part.startswith(title):
	# 				text_part = text_part.replace(title,'')
	# 			para = document.add_paragraph(text_part.strip())
	# 			article_dict[article_id]['text'].append(text_part.strip())
			# if file_open:
			# 	document.add_page_break()
			# 	unwanted = p.find('strong')
			# 	if unwanted:
			# 		unwanted.extract()
			# 		paragraph = p.text
			# 	if paragraph.strip() != '':
			# 		para = document.add_paragraph(paragraph.strip())
			# 		article_dict[article_id-1]['text'].append(paragraph.strip())
			# 	# f.close()
			# 	file_open = False
			
			
			# article_id+=1
			
			
			
			# f = open(f'../export_gf_old/{folder_name}/{date}_{article_id}.txt','w',encoding='utf8')
			
			# f.write(f'{title}\n')
			# file_open = True
			# texts = p.find_next_siblings('p')
		# if file_open:
		# 	paragraph = p.text
		# 	if paragraph.strip() != '' and paragraph.strip() != '[collapse]':
		# 		if paragraph == title:
		# 			continue
		# 		if paragraph.startswith(title):
		# 			paragraph = paragraph.replace(title,'')
		# 		# f.write(f'{paragraph.strip()}\n')
		# 		para = document.add_paragraph(paragraph.strip())
		# 		article_dict[article_id]['text'].append(paragraph.strip())
			
document.save(doc_name)

# sorting articles chronologically
sorted_articles = sorted(article_dict, key= lambda k: article_dict.get(k)['date_sort'])
sorted_articles = {k:article_dict[k] for k in sorted_articles}

# sorted(article_dict,key=lambda x: x['date_sort'])

sorted_doc_name = 'Spielberichte_Traktor_GF_old_sorted.docx'

sorted_document = Document()

sorted_document.add_heading('Spielberichte Traktor Großfeld alt sortiert', 0)


for article in sorted_articles:
	p = sorted_document.add_heading(sorted_articles[article]['heading'], level=1)
	run = p.add_run()
	run.add_break()
	for paragraph in sorted_articles[article]['text']:
		para = sorted_document.add_paragraph(paragraph)
	sorted_document.add_page_break()

	
sorted_document.save(sorted_doc_name)


			# with open(f'./export_gf_old/{date}_{article_id}.txt','w',encoding='utf8') as f:
			# 	for text in texts:
			# 		for child in text.children:
			# 			if child.name == 'strong' or child.contents[0].name == 'strong':
			# 				break
			# 		f.write(text.text)
					# if text.contents[0].name:
					# 	if text.contents[0].name == 'strong' or text.contents[1].name == 'strong':
					# 		break
					# else:
						# f.write(text.text)

				
			





		# text = p.text
		# try:
		# 	title = p.strong.text
		# 	title = traktor_clean_article(title)
		# 	try:
		# 		f.close()
		# 	except NameError:
		# 		pass
		# 	first_paragraph = True
			
		# except AttributeError as e:
		# 	title = None
		# 	first_paragraph = False
		# if first_paragraph:
		# 	text = traktor_clean_article(text)
		# 	# date = text.split('|')[0].strip()
		# 	try:
				
			
		# 	article_id+=1
		# 	f = open(f'./export_gf_old/{date}_{article_id}.txt','w',encoding='utf8')
		# 	f.write(f'{text}\n')
		# elif title is None:
		# 	try:
		# 		f.write(f'{text}\n')
		# 	except NameError:
		# 		pass
			
