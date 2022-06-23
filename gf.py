import requests
from bs4 import BeautifulSoup
from func import extract_date
from docx import Document
from docx.shared import Inches
from docx.text.run import Run
import operator

header = {
	'User-Agent' : 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:97.0) Gecko/20100101 Firefox/97.0'
}

doc_name = 'Spielberichte_Traktor_GF.docx'

document = Document()

document.add_heading('Spielberichte Traktor Großfeld', 0)


def get_articles(articles,word_doc):
	for article in articles:
		title = article.h2.text
		arcticle_id = article['id']
		print(f'parsing post {arcticle_id} with title {title}')
		try:
			date,new_title,date_sort = extract_date(title)
		except:
			date,new_title,date_sort = 'unknown',title,'0'
		heading = f'{date} | {new_title}'
		link = article.find('h2',class_='entry-title taggedlink').a['href']
		article_dict[arcticle_id] = {
			'date' : date,
			'title' : title,
			'link' : link,
			'date_sort' :  int(date_sort),
			'heading' : heading,
			'text' : []
		}
		url = link
		r = requests.get(url,headers=header).text
		soup = BeautifulSoup(r,'html.parser')
		paragraphs = soup.find_all('p')
		word_doc.add_heading(heading, level=1)
		# removed link from doc
		# document.add_paragraph(link, style='Intense Quote')
		for p in paragraphs:
			if p.text.strip() != '':
				para = word_doc.add_paragraph(p.text)
				article_dict[arcticle_id]['text'].append(p.text)
		word_doc.add_page_break()

		# articles in different text files:

		# with open(f'./export_gf/{date}_{arcticle_id}.txt','w',encoding='utf8') as f:
		# 	f.write(f'{title}\n')
		# 	f.write(f'{link}\n')
		# 	for p in paragraphs:
		# 		# print(p.text)
		# 		f.write(f'{p.text}\n')
		# 	f.close()
article_dict = {}

url = 'http://www.traktor-boxhagen.de/category/spielberichte-der-grossfeldmannschaft'

r = requests.get(url,headers=header).text

soup = BeautifulSoup(r,'html.parser')

articles = soup.find_all('article')
nav = soup.find('div',class_='nav-previous').a['href']
get_articles(articles,document)

while nav is not None:
	r = requests.get(nav,headers=header).text
	soup = BeautifulSoup(r,'html.parser')
	articles = soup.find_all('article')
	try:
		nav = soup.find('div',class_='nav-previous').a['href']
	except AttributeError as e:
		nav = None
	get_articles(articles,document)


# sorting articles chronologically
sorted_articles = sorted(article_dict, key= lambda k: article_dict.get(k)['date_sort'])
sorted_articles = {k:article_dict[k] for k in sorted_articles}

# sorted(article_dict,key=lambda x: x['date_sort'])

sorted_doc_name = 'Spielberichte_Traktor_GF_sorted.docx'

sorted_document = Document()

sorted_document.add_heading('Spielberichte Traktor Großfeld', 0)


for article in sorted_articles:
	# article_dict[arcticle_id] = {
	# 		'date' : date,
	# 		'title' : title,
	# 		'link' : link,
	# 		'date_sort' :  int(date_sort),
	# 		'heading' : heading,
	# 		'text' : []
	# 	}
	p = sorted_document.add_heading(sorted_articles[article]['heading'], level=1)
	run = p.add_run()
	run.add_break()
	for paragraph in sorted_articles[article]['text']:
		para = sorted_document.add_paragraph(paragraph)
	sorted_document.add_page_break()

	
document.save(doc_name)
sorted_document.save(sorted_doc_name)